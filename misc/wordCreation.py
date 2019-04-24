from docx import Document
from docx.enum.table import *
from docx.shared import *

document = Document()

table = document.add_table(rows=2, cols=6, style='Table Grid')
a = table.cell(0, 1)
b = table.cell(0, 5)
A = a.merge(b)



for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

for row in table.rows:
    row.height = Pt(20)

hdr_cells = table.rows[0].cells
psw_cells = table.rows[1].cells

newpassword = 'blablabla'

hdr_cells[0].paragraphs[0].text = 'Username'
hdr_cells[1].paragraphs[0].text ='qendrimvllasa@siemens.com'
psw_cells[0].paragraphs[0].text = 'Password'
psw_cells[1].paragraphs[0].text = newpassword
psw_cells[2].paragraphs[0].text = 'IoT Extension'
psw_cells[3].paragraphs[0].text = newpassword
psw_cells[4].paragraphs[0].text = 'Integration'
psw_cells[5].paragraphs[0].text = newpassword

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.bold = True

document.save('test' + '.docx')
