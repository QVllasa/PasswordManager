import io
import os
import datetime
import sys
import time
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.common.by import By

from PyQt5.QtCore import *
from PyQt5 import QtGui

from docx import Document
from docx.enum.table import *
from docx.shared import *
from qtpy.QtWidgets import *

from ui.AccDialog import Ui_AccDialog
from ui.Dialog import Ui_Dialog
from ui.mainwindow import Ui_MainWindow

app = QApplication(sys.argv)
app.setWindowIcon(QtGui.QIcon('mdsp_password_manager.ico'))


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.environ.get("_MEIPASS2", os.path.abspath("."))

    return os.path.join(base_path, relative_path)


class MainWindow(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)

        self.accounts = {}

        self.txt_data = resource_path("../accountLists.txt")

        try:
            with open(self.txt_data, 'r') as f:
                d = []
                for i in f:
                    if '=' in i:
                        d = i.strip().split(' ')
                        self.accounts[d[0]] = []
                    elif '=' not in i and not len(i.strip()) == 0:
                        i = i.strip()
                        self.accounts[d[0]].append(i.strip())
                    else:
                        continue
        except IOError:
            self.noList()

        self.currentBrowser = ''
        self.currentPassword = ''
        self.newPassword = ''

        self.nopassword_dialog = QMessageBox()
        self.finish_dialog = QMessageBox()
        self.error_dialog = QMessageBox()
        self.docx_dialog = QMessageBox()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        self.addCombo()
        browsers = ['Chrome', 'Firefox']
        self.ui.comboBox_2.addItems(browsers)
        self.ui.label_5.hide()
        self.ui.label_4.hide()
        self.ui.pushButton.clicked.connect(self.changePW)
        self.ui.addButton.clicked.connect(self.addList)
        self.ui.showButton.clicked.connect(self.showAcc)

    def addCombo(self):
        for key, value in self.accounts.items():
            self.ui.comboBox.addItem(key)

    def noList(self):
        self.noList_dialog = QMessageBox()
        self.noList_dialog.setIcon(QMessageBox.Critical)
        self.noList_dialog.setText('No accountList.txt found!')
        self.noList_dialog.show()

    def changePW(self):
        self.currentBrowser = self.ui.comboBox_2.currentText()
        self.currentPassword = self.ui.inputCurrent.text()
        self.newPassword = self.ui.inputNew.text()
        self.account = self.ui.comboBox.currentText()
        self.state = str(self.ui.checkBox.checkState())

        if not self.currentPassword == '' and not self.newPassword == '':
            self.obj = Worker(self.account, self.currentBrowser, self.currentPassword, self.newPassword, self.accounts,
                              self.state)
            self.obj.message.connect(self.errorDialog)
            self.obj.finished.connect(self.done)
            self.obj.progress.connect(self.loadingBar)
            self.obj.docFinish.connect(self.docxDialog)
            self.obj.label4.connect(self.okay4)
            self.obj.label5.connect(self.okay5)
            self.obj.start()
        else:
            self.nopassword()

    def addList(self):
        self.Adddialog = QDialog()
        self.uiAddDialog = Ui_Dialog()
        self.uiAddDialog.setupUi(self.Adddialog)
        self.Adddialog.show()
        self.uiAddDialog.addRow.clicked.connect(self.addCell)
        rsp = self.Adddialog.exec_()
        if rsp == QDialog.Accepted:
            listname = self.uiAddDialog.listName.text()
            if listname:
                if not listname in self.accounts:
                    self.accounts[listname] = []
                    with open(self.txt_data, 'a') as f:
                        f.write('\n' + listname + ' ' + '=\n')

                    print(self.accounts)
            for i in range(0, self.uiAddDialog.accTable.rowCount()):
                if not self.uiAddDialog.accTable.item(i, 0) == None:
                    s = self.uiAddDialog.accTable.item(i, 0).text()
                    with open(self.txt_data, 'a') as f:
                        f.write(s + '\n')


                else:
                    print('leer')


        else:
            print('Cancel')

    def addCell(self):
        row = self.uiAddDialog.accTable.rowCount()
        self.uiAddDialog.accTable.insertRow(row)

    def showAcc(self):
        self.Accdialog = QDialog()
        self.uiAccDialog = Ui_AccDialog()
        self.uiAccDialog.setupUi(self.Accdialog)
        for key, value in self.accounts.items():
            self.uiAccDialog.comboBox.addItem(key)
        if self.uiAccDialog.comboBox.currentText() in self.accounts:
            self.on_comboBox_changed(self.uiAccDialog.comboBox.currentText())
        self.uiAccDialog.comboBox.currentTextChanged.connect(self.on_comboBox_changed)
        self.Accdialog.show()
        self.Accdialog.exec_()

    def on_comboBox_changed(self, key):
        while self.uiAccDialog.tableWidget.rowCount() > 0:
            self.uiAccDialog.tableWidget.removeRow(0)
        if self.uiAccDialog.comboBox.currentText() == key:
            for i, j in self.accounts.items():
                if key == i:
                    for s in j:
                        row = self.uiAccDialog.tableWidget.rowCount()
                        self.uiAccDialog.tableWidget.insertRow(row)
                        self.uiAccDialog.tableWidget.setItem(row, 0, QTableWidgetItem(str(s)))

    def errorDialog(self, errorText):
        print(errorText)
        self.error_dialog.setIcon(QMessageBox.Critical)
        self.error_dialog.setText('Problem with password in account:  ' + errorText)
        self.error_dialog.show()

    def nopassword(self):
        self.nopassword_dialog.setIcon(QMessageBox.Information)
        self.nopassword_dialog.setText('Please fill out emtpy spaces!')
        self.nopassword_dialog.show()

    def done(self, str):
        self.finish_dialog.setIcon(QMessageBox.Information)
        self.finish_dialog.setText('Passwords changed to:   ' + str)
        self.finish_dialog.show()

    def loadingBar(self, percentage):
        print(percentage)
        self.ui.progressBar.setValue(percentage)

    def docxDialog(self, str):
        self.docx_dialog.setIcon(QMessageBox.Information)
        self.docx_dialog.setText(str)
        self.docx_dialog.show()

    def okay4(self, okay1):
        self.ui.label_4.setText(okay1)
        self.ui.label_4.show()

    def okay5(self, okay2):
        self.ui.label_5.setText(okay2)
        self.ui.label_5.show()


class Worker(QThread):
    message = pyqtSignal(str)
    finished = pyqtSignal(str)
    progress = pyqtSignal(float)
    label4 = pyqtSignal(str)
    label5 = pyqtSignal(str)
    docFinish = pyqtSignal(str)

    def __init__(self, account, brwser, curPass, newPass, accountList, state):
        QThread.__init__(self)

        self.currentBrowser = brwser
        self.currentPassword = curPass
        self.newPassword = newPass
        self.account = account
        self.accountList = accountList
        self.state = state

    def run(self):
        document = Document()
        document.sections[0].left_margin = Mm(10)
        document.sections[0].right_margin = Mm(10)

        count = float(0)
        self.progress.emit(count)
        while round(count) < float(100):
            for acc, address in self.accountList.items():
                if self.account == acc:
                    print(acc)
                    for user in self.accountList[acc]:

                        table = document.add_table(rows=2, cols=6, style='Table Grid')
                        for row in table.rows:
                            for cell in row.cells:
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

                        for row in table.rows:
                            row.height = Pt(20)

                        hdr_cells = table.rows[0].cells
                        psw_cells = table.rows[1].cells

                        hdr_cells[0].paragraphs[0].text = 'Username'
                        psw_cells[0].paragraphs[0].text = 'Password'

                        if 'connectivity' in user:
                            a = table.cell(0, 1)
                            b = table.cell(0, 5)
                            A = a.merge(b)
                            psw_cells[2].paragraphs[0].text = 'IoT Extension'
                            psw_cells[4].paragraphs[0].text = 'MC Integration'

                        else:
                            a = table.cell(0, 1)
                            b = table.cell(0, 5)
                            c = table.cell(1, 1)
                            d = table.cell(1, 5)
                            A = a.merge(b)
                            B = c.merge(d)

                        print(acc)
                        b = len(self.accountList[acc])
                        # print(b)
                        count += float(100) / float(b)
                        print(count)
                        print(user)
                        page = "https://www2.industrysoftware.automation.siemens.com/webkey/"

                        if self.currentBrowser == 'Chrome':
                            optionsChrome = webdriver.ChromeOptions()
                            if self.state == '2':
                                optionsChrome.add_argument("--window-size=1920,1080")
                                optionsChrome.add_argument('--start-maximized')
                                optionsChrome.headless = True
                            driver = webdriver.Chrome(
                                executable_path=resource_path('webdriver/windows/chromedriver.exe'),
                                options=optionsChrome)

                        if self.currentBrowser == 'Firefox':
                            optionsFirefox = webdriver.FirefoxOptions()
                            if self.state == '2':
                                optionsFirefox.headless = True
                            driver = webdriver.Firefox(
                                executable_path=resource_path('webdriver/windows/geckodriver.exe'),
                                options=optionsFirefox,
                            )

                        driver.implicitly_wait(3)
                        # driver.minimize_window()
                        driver.get(page)
                        driver.find_element(By.XPATH, "//tr[5]/td[2]/ul/font/li/a/font").click()
                        driver.find_element(By.XPATH, "//td[2]/input").click()
                        driver.find_element(By.NAME, "WebKey_Username").send_keys(user)
                        driver.find_element(By.NAME, "Existing_WebKey_Password").send_keys(self.currentPassword)
                        driver.find_element(By.NAME, "pass").click()
                        driver.find_element(By.NAME, "pass").send_keys(self.newPassword)
                        driver.find_element(By.NAME, "repass").click()
                        driver.find_element(By.NAME, "repass").send_keys(self.newPassword)
                        driver.find_element(By.XPATH, "//div[3]/div[2]/div/form/fieldset/input").click()
                        time.sleep(1)

                        # wrong password entered
                        try:
                            driver.find_element(By.XPATH, "//h2[contains(.,'WebKey Error')]")
                            self.message.emit(user)
                            self.progress.emit(count)
                            hdr_cells[1].paragraphs[0].text = user
                            psw_cells[1].paragraphs[0].text = 'Wrong PW!'

                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = 'Calibri'
                                            run.font.size = Pt(12)
                                            run.bold = True
                            document.add_paragraph('')
                            driver.quit()
                            continue
                        except NoSuchElementException:
                            pass

                        # new password matches older one
                        try:
                            driver.find_element(By.XPATH,
                                                "//h2[contains(.,'The following message was returned from the WebKey Server:')]")
                            self.message.emit(user)
                            self.progress.emit(count)
                            hdr_cells[1].paragraphs[0].text = user
                            psw_cells[1].paragraphs[0].text = 'New Password matches older one or to many attempts!!'
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = 'Calibri'
                                            run.font.size = Pt(12)
                                            run.bold = True
                            document.add_paragraph('')
                            driver.quit()
                            continue
                        except NoSuchElementException:
                            pass

                        # Your Password has been Changed
                        try:
                            driver.find_element(By.XPATH, "//h3[contains(.,'Your Password has been Changed.')]")

                            print('found')
                            self.progress.emit(count)
                            self.label4.emit('OK')
                            self.label5.emit('OK')
                            hdr_cells[1].paragraphs[0].text = user
                            psw_cells[1].paragraphs[0].text = self.newPassword
                            psw_cells[3].paragraphs[0].text = self.newPassword
                            psw_cells[5].paragraphs[0].text = self.newPassword
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = 'Calibri'
                                            run.font.size = Pt(12)
                                            run.bold = True

                            document.add_paragraph('')
                            driver.quit()
                            print(str(count) + '%')
                        except NoSuchElementException:
                            print('not found')
                            pass

        self.finished.emit(self.newPassword)
        document.add_paragraph(str(datetime.datetime.now()))
        for acc in self.accountList:
            if acc == self.account:
                document.save(
                    resource_path("../" + acc + '_' + 'CW' + str(datetime.datetime.today().isocalendar()[1]) + '.docx'))
        self.docFinish.emit('Word file with all changed accounts created and saved in application folder! :)')


window = MainWindow()
window.show()
window.setWindowIcon(QtGui.QIcon('mdsp_password_manager.ico'))


sys.exit(app.exec_())
